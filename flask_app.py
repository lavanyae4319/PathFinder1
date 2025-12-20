"""
Flask Web Application - Intelligent Resume Analyzer & Job Match Prediction System
Features: Resume Analysis, Job Prediction, Resume Generator
"""

from flask import Flask, render_template, request, jsonify, send_file, flash, redirect, url_for, session
from werkzeug.utils import secure_filename
import os
import json
from datetime import datetime
import secrets
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from simple_resume_parser import SimpleResumeParser
from job_predictor import JobPredictor
from gemini_analyzer import GeminiResumeAnalyzer
from simple_analyzer import analyze_resume_simple
from config import (
    RESUME_UPLOAD_FOLDER, RESULTS_FOLDER, SUPPORTED_RESUME_FORMATS,
    MAX_RESUME_SIZE_MB
)

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = secrets.token_hex(16)
app.config['UPLOAD_FOLDER'] = RESUME_UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_RESUME_SIZE_MB * 1024 * 1024  # MB to bytes

# Initialize components
try:
    resume_parser = SimpleResumeParser()
    job_predictor = JobPredictor()
    gemini_analyzer = GeminiResumeAnalyzer()  # For comprehensive analysis
    print("✅ Application components initialized successfully!")
    print("✅ Using Simple Resume Parser (No API Key Required)")
    if gemini_analyzer.client:
        print("✅ Gemini 2.5 Flash available for comprehensive analysis")
except Exception as e:
    print(f"⚠ Warning: Could not initialize components: {str(e)}")
    resume_parser = None
    job_predictor = None
    gemini_analyzer = None


# Helper Functions
def allowed_file(filename):
    """Check if uploaded file is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in [f.replace('.', '') for f in SUPPORTED_RESUME_FORMATS]


def format_experience_years(exp):
    """Format experience years from parsed resume"""
    if isinstance(exp, (int, float)):
        return float(exp)
    elif isinstance(exp, str):
        import re
        numbers = re.findall(r'\d+\.?\d*', exp)
        return float(numbers[0]) if numbers else 0
    return 0


# Routes
@app.route('/')
def index():
    """Home page"""
    return render_template('index.html')


@app.route('/analyze')
def analyze_page():
    """Resume Analysis page"""
    return render_template('analyze.html')


@app.route('/generator')
def generator_page():
    """Resume Generator page"""
    return render_template('generator.html')


@app.route('/about')
def about_page():
    """About page"""
    return render_template('about.html')


@app.route('/job-matcher')
def job_matcher_page():
    """Job Matcher page"""
    return render_template('job_matcher.html')


@app.route('/api/analyze', methods=['POST'])
def analyze_resume():
    """API endpoint for comprehensive resume analysis using Gemini AI"""
    try:
        # Check if file was uploaded
        if 'resume' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['resume']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': f'Invalid file format. Supported: {", ".join(SUPPORTED_RESUME_FORMATS)}'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)
        
        # Get top_k from request
        top_k = int(request.form.get('top_k', 10))
        
        # Comprehensive analysis using Gemini AI
        print(f"🤖 Analyzing resume with Gemini 2.5 Flash: {filename}")
        gemini_analysis = None
        if gemini_analyzer and gemini_analyzer.client:
            try:
                gemini_analysis = gemini_analyzer.analyze_resume_comprehensive(filepath)
                print("✅ Gemini analysis completed")
            except Exception as e:
                print(f"⚠ Gemini analysis failed: {e}, using fallback")
        
        # Fallback to simple parser if Gemini not available
        if not gemini_analysis:
            parsed_resume = resume_parser.parse_resume(filepath)
            simple_analysis = analyze_resume_simple(parsed_resume)
            gemini_analysis = {
                'parsed_data': parsed_resume,
                'analysis': simple_analysis,
                'meta': {
                    'provider': 'simple',
                    'model': 'rule-based'
                }
            }
        
        # Get ML-based job recommendations (using trained model)
        print(f"🎯 Finding matching jobs using ML model...")
        parsed_for_ml = gemini_analysis.get('parsed_data', {})
        # Build ML input from either Gemini or SimpleResumeParser formats
        if isinstance(parsed_for_ml, dict) and 'personal_info' in parsed_for_ml:
            ml_format = {
                'Personal Information': parsed_for_ml.get('personal_info', {}),
                'Professional Summary': parsed_for_ml.get('professional_summary', ''),
                'Skills': parsed_for_ml.get('skills', []),
                'Work Experience': parsed_for_ml.get('work_experience', []),
                'Education': parsed_for_ml.get('education', []),
                'Total Years of Experience': parsed_for_ml.get('total_experience_years', 0)
            }
        else:
            ml_format = {
                'Personal Information': {
                    'Name': parsed_for_ml.get('name', 'N/A'),
                    'Email': parsed_for_ml.get('email', 'N/A'),
                    'Phone': parsed_for_ml.get('phone', 'N/A'),
                    'Location': parsed_for_ml.get('location', 'N/A')
                },
                'Professional Summary': parsed_for_ml.get('professional_summary', parsed_for_ml.get('summary', '')),
                'Skills': parsed_for_ml.get('skills', []),
                'Work Experience': parsed_for_ml.get('work_experience', parsed_for_ml.get('experience', [])),
                'Education': parsed_for_ml.get('education', []),
                'Total Years of Experience': parsed_for_ml.get('total_experience_years', parsed_for_ml.get('experience_years', 0))
            }
        
        recommendations = job_predictor.get_job_recommendations(ml_format, top_k=top_k)
        
        # Prepare comprehensive response
        parsed_info = gemini_analysis.get('parsed_data', {})
        # Support both Gemini format (with personal_info) and SimpleResumeParser format (flat fields)
        if isinstance(parsed_info, dict) and 'personal_info' in parsed_info:
            personal_info = parsed_info.get('personal_info', {})
            professional_summary = parsed_info.get('professional_summary', 'N/A')
            experience_years = parsed_info.get('total_experience_years', 0)
            skills = parsed_info.get('skills', [])
            education = parsed_info.get('education', [])
            work_experience = parsed_info.get('work_experience', [])
            certifications = parsed_info.get('certifications', [])
            preferred_job_titles = parsed_info.get('preferred_job_titles', [])
            preferred_industries = parsed_info.get('preferred_industries', [])
        else:
            # SimpleResumeParser returns flattened fields
            personal_info = {
                'name': parsed_info.get('name', 'N/A'),
                'email': parsed_info.get('email', 'N/A'),
                'phone': parsed_info.get('phone', 'N/A'),
                'location': parsed_info.get('location', 'N/A'),
                'linkedin': parsed_info.get('linkedin', 'N/A')
            }
            professional_summary = parsed_info.get('professional_summary', parsed_info.get('summary', 'N/A'))
            # Accept both keys for experience years
            experience_years = parsed_info.get('total_experience_years', parsed_info.get('experience_years', 0))
            skills = parsed_info.get('skills', [])
            education = parsed_info.get('education', [])
            work_experience = parsed_info.get('work_experience', parsed_info.get('experience', []))
            certifications = parsed_info.get('certifications', [])
            preferred_job_titles = parsed_info.get('preferred_job_titles', [])
            preferred_industries = parsed_info.get('preferred_industries', [])

        analysis = gemini_analysis.get('analysis', {})
        
        response_data = {
            'success': True,
            'filename': filename,
            'ai_meta': gemini_analysis.get('meta', {}),
            'parsed_data': {
                'name': personal_info.get('name', 'N/A'),
                'email': personal_info.get('email', 'N/A'),
                'phone': personal_info.get('phone', 'N/A'),
                'location': personal_info.get('location', 'N/A'),
                'linkedin': personal_info.get('linkedin', 'N/A'),
                'experience_years': experience_years,
                'summary': professional_summary,
                'skills': skills,
                'education': education,
                'work_experience': work_experience,
                'certifications': certifications,
                'preferred_job_titles': preferred_job_titles,
                'preferred_industries': preferred_industries
            },
            'ai_analysis': {
                'overall_score': analysis.get('overall_score', 0),
                'strengths': analysis.get('strengths', []),
                'weaknesses': analysis.get('weaknesses', []),
                'suggestions': analysis.get('suggestions', []),
                'job_recommendations': analysis.get('job_recommendations', []),
                'skills_gap_analysis': analysis.get('skills_gap_analysis', []),
                'career_path_suggestions': analysis.get('career_path_suggestions', []),
                'ats_optimization': analysis.get('ats_optimization', []),
                'keywords_to_add': analysis.get('keywords_to_add', []),
                'summary_paragraph': analysis.get('summary_paragraph', '')
            },
            'ml_predictions': {
                'predicted_category': recommendations['candidate_summary']['predicted_category'],
                'confidence': round(recommendations['candidate_summary']['category_confidence'] * 100, 2),
                'total_matches': recommendations['total_matches'],
                'avg_match_score': round(recommendations['avg_match_score'], 2)
            },
            'matching_jobs': [
                {
                    'Job_Title': job.get('jobtitle', 'N/A'),
                    'Company': job.get('company', 'N/A'),
                    'Location': job.get('primary_location', 'N/A'),
                    'Job_Category': job.get('job_category', 'N/A'),
                    'Job_Experience_Required': job.get('experience', 'N/A'),
                    'Skills': job.get('skills', ''),
                    'Job_Description': job.get('jobdescription', ''),
                    'match_score': job.get('similarity_score', 0),
                    'match_percentage': job.get('match_percentage', 0)
                }
                for job in recommendations['matching_jobs'][:top_k]
            ],
            'top_companies': recommendations.get('top_companies', {}),
            'top_locations': recommendations.get('top_locations', {})
        }
        
        # Save results
        result_filename = f"analysis_{timestamp}.json"
        result_path = os.path.join(RESULTS_FOLDER, result_filename)
        with open(result_path, 'w', encoding='utf-8') as f:
            json.dump(response_data, f, indent=2)
        
        print(f"✅ Analysis complete. Overall score: {analysis.get('overall_score', 'N/A')}")
        
        return jsonify(response_data)
    
    except Exception as e:
        print(f"Error during analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/job-match', methods=['POST'])
def job_match():
    """API endpoint for job matching without full resume parsing"""
    try:
        # Check if file was uploaded
        if 'resume' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['resume']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': f'Invalid file format. Supported: {", ".join(SUPPORTED_RESUME_FORMATS)}'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)
        
        # Get top_k from request
        top_k = int(request.form.get('top_k', 10))
        
        # Parse resume using AI
        print(f"Parsing resume for job matching: {filename}")
        parsed_resume = resume_parser.parse_resume(filepath)
        
        # Get job recommendations
        print(f"Finding matching jobs with ML model...")
        recommendations = job_predictor.get_job_recommendations(parsed_resume, top_k=top_k)
        
        # Transform matching jobs into the format expected by job_matcher.html
        transformed_jobs = []
        for job in recommendations['matching_jobs'][:top_k]:
            transformed_jobs.append({
                # Original fields kept (if frontend ever needs them)
                **job,
                # Fields expected by job_matcher.html JavaScript
                'Job_Title': job.get('jobtitle') or job.get('job_title') or job.get('jobTitle') or 'Position',
                'Match_Score': float(job.get('match_percentage', 0)) / 100.0
                    if job.get('match_percentage') is not None else 0.0,
                'Company': job.get('company', 'Company'),
                'Location': job.get('primary_location') or job.get('location') or 'N/A',
                'Job_Category': job.get('job_category', 'N/A'),
                'Job_Description': job.get('jobdescription') or job.get('job_description') or '',
                'Job_Experience_Required': job.get('experience') or job.get('experience_required') or 'N/A',
            })

        # Prepare response
        response_data = {
            'success': True,
            'filename': filename,
            'candidate': {
                'name': parsed_resume.get('Personal Information', {}).get('Name', 'N/A'),
                'email': parsed_resume.get('Personal Information', {}).get('Email', 'N/A'),
                'phone': parsed_resume.get('Personal Information', {}).get('Phone', 'N/A'),
                'experience_years': format_experience_years(
                    parsed_resume.get('Total Years of Experience', 0)
                ),
                'skills': parsed_resume.get('Skills', []),
                'predicted_category': recommendations['candidate_summary']['predicted_category'],
                'confidence': round(
                    recommendations['candidate_summary']['category_confidence'] * 100, 2
                ),
            },
            'match_summary': {
                'total_matches': recommendations['total_matches'],
                'avg_match_score': round(recommendations['avg_match_score'], 2),
            },
            'matching_jobs': transformed_jobs,
            'top_companies': recommendations.get('top_companies', {}),
            'top_locations': recommendations.get('top_locations', {}),
        }
        
        return jsonify(response_data)
    
    except Exception as e:
        print(f"Error during job matching: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-resume', methods=['POST'])
def generate_resume():
    """API endpoint for resume generation"""
    try:
        data = request.json
        
        # Extract data
        personal_info = data.get('personal_info', {})
        summary = data.get('summary', '')
        skills = data.get('skills', [])
        experience = data.get('experience', [])
        education = data.get('education', [])
        certifications = data.get('certifications', [])
        format_type = data.get('format', 'docx')  # docx or pdf
        
        # Generate resume based on format
        if format_type == 'pdf':
            file_buffer = generate_pdf_resume(personal_info, summary, skills, experience, education, certifications)
            filename = f"resume_{personal_info.get('name', 'generated').replace(' ', '_')}.pdf"
            mimetype = 'application/pdf'
        else:  # docx
            file_buffer = generate_docx_resume(personal_info, summary, skills, experience, education, certifications)
            filename = f"resume_{personal_info.get('name', 'generated').replace(' ', '_')}.docx"
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        return send_file(
            file_buffer,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        print(f"Error generating resume: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


def generate_docx_resume(personal_info, summary, skills, experience, education, certifications):
    """Generate resume in DOCX format"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - Name
    name = doc.add_heading(personal_info.get('name', 'Your Name'), 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.runs[0]
    name_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Contact Information
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = []
    if personal_info.get('email'):
        contact_text.append(personal_info['email'])
    if personal_info.get('phone'):
        contact_text.append(personal_info['phone'])
    if personal_info.get('location'):
        contact_text.append(personal_info['location'])
    if personal_info.get('linkedin'):
        contact_text.append(personal_info['linkedin'])
    
    contact_para.add_run(' | '.join(contact_text))
    contact_para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Professional Summary
    if summary:
        doc.add_heading('Professional Summary', 2).runs[0].font.color.rgb = RGBColor(0, 51, 102)
        doc.add_paragraph(summary)
        doc.add_paragraph()
    
    # Skills
    if skills:
        doc.add_heading('Skills', 2).runs[0].font.color.rgb = RGBColor(0, 51, 102)
        skills_para = doc.add_paragraph()
        skills_para.add_run(', '.join(skills))
        doc.add_paragraph()
    
    # Work Experience
    if experience:
        doc.add_heading('Work Experience', 2).runs[0].font.color.rgb = RGBColor(0, 51, 102)
        for exp in experience:
            # Job title and company
            job_para = doc.add_paragraph()
            job_run = job_para.add_run(f"{exp.get('title', 'Position')} at {exp.get('company', 'Company')}")
            job_run.bold = True
            
            # Duration
            if exp.get('duration'):
                duration_para = doc.add_paragraph(exp['duration'])
                duration_para.runs[0].italic = True
            
            # Responsibilities
            if exp.get('responsibilities'):
                for resp in exp['responsibilities']:
                    doc.add_paragraph(resp, style='List Bullet')
            
            doc.add_paragraph()
    
    # Education
    if education:
        doc.add_heading('Education', 2).runs[0].font.color.rgb = RGBColor(0, 51, 102)
        for edu in education:
            edu_para = doc.add_paragraph()
            edu_text = f"{edu.get('degree', 'Degree')} in {edu.get('field', 'Field')}"
            if edu.get('institution'):
                edu_text += f" - {edu['institution']}"
            if edu.get('year'):
                edu_text += f" ({edu['year']})"
            edu_para.add_run(edu_text).bold = True
            
            if edu.get('gpa'):
                doc.add_paragraph(f"GPA: {edu['gpa']}")
            doc.add_paragraph()
    
    # Certifications
    if certifications:
        doc.add_heading('Certifications', 2).runs[0].font.color.rgb = RGBColor(0, 51, 102)
        for cert in certifications:
            doc.add_paragraph(cert, style='List Bullet')
    
    # Save to BytesIO
    file_buffer = BytesIO()
    doc.save(file_buffer)
    file_buffer.seek(0)
    
    return file_buffer


def generate_pdf_resume(personal_info, summary, skills, experience, education, certifications):
    """Generate resume in PDF format"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    # Container for elements
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#003366'),
        spaceAfter=12,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#003366'),
        spaceAfter=6,
        borderWidth=1,
        borderColor=colors.HexColor('#003366'),
        borderPadding=3
    )
    
    # Name
    elements.append(Paragraph(personal_info.get('name', 'Your Name'), title_style))
    
    # Contact
    contact_text = []
    if personal_info.get('email'):
        contact_text.append(personal_info['email'])
    if personal_info.get('phone'):
        contact_text.append(personal_info['phone'])
    if personal_info.get('location'):
        contact_text.append(personal_info['location'])
    
    contact_para = Paragraph(' | '.join(contact_text), styles['Normal'])
    contact_para.alignment = 1
    elements.append(contact_para)
    elements.append(Spacer(1, 0.2*inch))
    
    # Summary
    if summary:
        elements.append(Paragraph('Professional Summary', heading_style))
        elements.append(Paragraph(summary, styles['Normal']))
        elements.append(Spacer(1, 0.15*inch))
    
    # Skills
    if skills:
        elements.append(Paragraph('Skills', heading_style))
        elements.append(Paragraph(', '.join(skills), styles['Normal']))
        elements.append(Spacer(1, 0.15*inch))
    
    # Experience
    if experience:
        elements.append(Paragraph('Work Experience', heading_style))
        for exp in experience:
            job_text = f"<b>{exp.get('title', 'Position')} at {exp.get('company', 'Company')}</b>"
            elements.append(Paragraph(job_text, styles['Normal']))
            
            if exp.get('duration'):
                elements.append(Paragraph(f"<i>{exp['duration']}</i>", styles['Normal']))
            
            if exp.get('responsibilities'):
                for resp in exp['responsibilities']:
                    elements.append(Paragraph(f"• {resp}", styles['Normal']))
            
            elements.append(Spacer(1, 0.1*inch))
    
    # Education
    if education:
        elements.append(Paragraph('Education', heading_style))
        for edu in education:
            edu_text = f"<b>{edu.get('degree', 'Degree')} in {edu.get('field', 'Field')}</b>"
            if edu.get('institution'):
                edu_text += f" - {edu['institution']}"
            if edu.get('year'):
                edu_text += f" ({edu['year']})"
            elements.append(Paragraph(edu_text, styles['Normal']))
            elements.append(Spacer(1, 0.1*inch))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    return buffer


@app.route('/api/system-status')
def system_status():
    """Check system status"""
    status = {
        'parser_ready': resume_parser is not None,
        'predictor_ready': job_predictor is not None,
        'parser_type': 'Simple Text Extraction',
        'total_jobs': len(job_predictor.job_data) if job_predictor and job_predictor.job_data is not None else 0
    }
    return jsonify(status)


@app.errorhandler(413)
def too_large(e):
    """Handle file too large error"""
    return jsonify({'error': f'File too large. Maximum size: {MAX_RESUME_SIZE_MB}MB'}), 413


@app.errorhandler(404)
def not_found(e):
    """Handle 404 errors"""
    return render_template('404.html'), 404


@app.errorhandler(500)
def server_error(e):
    """Handle 500 errors"""
    return render_template('500.html'), 500


if __name__ == '__main__':
    # Create necessary directories
    os.makedirs(RESUME_UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(RESULTS_FOLDER, exist_ok=True)
    os.makedirs('static', exist_ok=True)
    os.makedirs('templates', exist_ok=True)
    
    print("\n" + "="*70)
    print("🚀 PATHFINDER - RESUME ANALYZER AND JOB PREDICTOR")
    print("="*70)
    print(f"\n✓ Resume Parser: Simple Text Extraction (No API Required)")
    print(f"✓ Job Matching: Trained ML Model (Gradient Boosting)")
    print(f"✓ Upload Folder: {RESUME_UPLOAD_FOLDER}")
    print(f"✓ Results Folder: {RESULTS_FOLDER}")
    print(f"\n🌐 Starting Flask server...")
    print("="*70)
    
    # Run the app with debug mode for live reload
    app.run(debug=True, host='0.0.0.0', port=5000)
