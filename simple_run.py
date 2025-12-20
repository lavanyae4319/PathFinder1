"""
Flask App with ML Model - Job Prediction using Trained Model
"""
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import os
import json
from datetime import datetime
from simple_resume_parser import SimpleResumeParser
from job_predictor import JobPredictor
from simple_analyzer import analyze_resume_simple

app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY') or 'please-set-flask-secret-key'
app.config['UPLOAD_FOLDER'] = 'uploads/resumes'
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB

# Create upload folder
os.makedirs('uploads/resumes', exist_ok=True)
os.makedirs('results', exist_ok=True)

# Initialize parsers and predictors
print("Loading ML models...")
resume_parser = SimpleResumeParser()
job_predictor = JobPredictor()
print("[OK] ML models loaded successfully!")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ['pdf', 'docx', 'txt']

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze')
def analyze_page():
    return render_template('analyze.html')

@app.route('/generator')
def generator_page():
    return render_template('generator.html')

@app.route('/job-matcher')
def job_matcher_page():
    return render_template('job_matcher.html')

@app.route('/about')
def about_page():
    return render_template('about.html')

@app.route('/test-form')
def test_form_page():
    return render_template('test_form.html')

@app.route('/api/analyze', methods=['POST'])
def analyze_resume():
    """Analyze resume using trained ML model"""
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file format. Supported: PDF, DOCX, TXT'}), 400
        
        # Save file
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)
        
        top_k = int(request.form.get('top_k', 10))
        
        # Parse resume (simple text extraction)
        print(f"Parsing resume: {filename}")
        parsed_resume = resume_parser.parse_resume(filepath)
        
        # Analyze resume for strengths, weaknesses, suggestions
        print(f"Analyzing resume quality...")
        simple_analysis = analyze_resume_simple(parsed_resume)
        
        # Get job predictions using trained ML model
        print(f"Predicting job category and finding matches...")
        recommendations = job_predictor.get_job_recommendations(parsed_resume, top_k=top_k)
        
        # Prepare response
        response_data = {
            'success': True,
            'filename': filename,
            'ai_meta': {
                'provider': 'simple',
                'model': 'rule-based'
            },
            'parsed_data': {
                'name': parsed_resume['Personal Information']['Name'],
                'email': parsed_resume['Personal Information']['Email'],
                'phone': parsed_resume['Personal Information']['Phone'],
                'location': parsed_resume['Personal Information']['Location'],
                'experience_years': parsed_resume['Total Years of Experience'],
                'summary': parsed_resume['Professional Summary'],
                'skills': parsed_resume['Skills'],
                'education': parsed_resume['Education'],
                'work_experience': parsed_resume['Work Experience']
            },
            'ai_analysis': simple_analysis,
            'predictions': {
                'predicted_category': recommendations['candidate_summary']['predicted_category'],
                'confidence': round(recommendations['candidate_summary']['category_confidence'] * 100, 2),
                'total_matches': recommendations['total_matches'],
                'avg_match_score': round(recommendations['avg_match_score'], 2)
            },
            'matching_jobs': recommendations['matching_jobs'][:top_k],
            'top_companies': recommendations.get('top_companies', {}),
            'top_locations': recommendations.get('top_locations', {})
        }
        
        # Save results
        result_filename = f"analysis_{timestamp}.json"
        result_path = os.path.join('results', result_filename)
        with open(result_path, 'w', encoding='utf-8') as f:
            json.dump(response_data, f, indent=2)
        
        return jsonify(response_data)
    
    except Exception as e:
        print(f"Error during analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/job-match', methods=['POST'])
def job_match():
    """Job matching using trained ML model"""
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file format. Supported: PDF, DOCX, TXT'}), 400
        
        # Save file
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)
        
        top_k = int(request.form.get('top_k', 10))
        
        # Parse resume
        print(f"[INFO] Parsing resume: {filename}")
        parsed_resume = resume_parser.parse_resume(filepath)
        
        # Get job predictions using trained ML model
        print(f"[INFO] ML Model: Predicting job category and finding matches...")
        recommendations = job_predictor.get_job_recommendations(parsed_resume, top_k=top_k)
        
        # Prepare response
        response_data = {
            'success': True,
            'filename': filename,
            'parsed_data': {
                'name': parsed_resume['Personal Information']['Name'],
                'email': parsed_resume['Personal Information']['Email'],
                'phone': parsed_resume['Personal Information']['Phone'],
                'location': parsed_resume['Personal Information']['Location'],
                'experience_years': f"{parsed_resume['Total Years of Experience']} years",
                'summary': parsed_resume['Professional Summary'],
                'skills': parsed_resume['Skills'],
                'education': parsed_resume['Education'],
                'work_experience': parsed_resume['Work Experience']
            },
            'predictions': {
                'predicted_category': recommendations['candidate_summary']['predicted_category'],
                'confidence': round(recommendations['candidate_summary']['category_confidence'] * 100, 2),
                'total_matches': recommendations['total_matches'],
                'avg_match_score': round(recommendations['avg_match_score'], 2)
            },
            'matching_jobs': recommendations['matching_jobs'][:top_k],
            'top_companies': recommendations.get('top_companies', {}),
            'top_locations': recommendations.get('top_locations', {})
        }
        
        print(f"[OK] Found {response_data['predictions']['total_matches']} matches")
        print(f"[OK] Predicted Category: {response_data['predictions']['predicted_category']}")
        
        return jsonify(response_data)
    
    except Exception as e:
        print(f"[ERROR] Job matching failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-resume', methods=['POST'])
def generate_resume():
    """API endpoint for resume generation"""
    try:
        from io import BytesIO
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib import colors
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        data = request.json
        
        # Extract data
        personal_info = data.get('personal_info', {})
        summary = data.get('summary', '')
        skills = data.get('skills', [])
        experience = data.get('experience', [])
        education = data.get('education', [])
        certifications = data.get('certifications', [])
        format_type = data.get('format', 'docx')  # docx or pdf
        
        # Generate resume based on format
        if format_type == 'pdf':
            file_buffer = generate_pdf_resume(personal_info, summary, skills, experience, education, certifications)
            filename = f"resume_{personal_info.get('name', 'generated').replace(' ', '_')}.pdf"
            mimetype = 'application/pdf'
        else:  # docx
            file_buffer = generate_docx_resume(personal_info, summary, skills, experience, education, certifications)
            filename = f"resume_{personal_info.get('name', 'generated').replace(' ', '_')}.docx"
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        return send_file(
            file_buffer,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        print(f"Error generating resume: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def generate_docx_resume(personal_info, summary, skills, experience, education, certifications):
    """Generate resume in DOCX format"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - Name
    name = doc.add_heading(personal_info.get('name', 'Your Name'), 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.runs[0]
    name_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Contact Information
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = []
    if personal_info.get('email'):
        contact_text.append(personal_info['email'])
    if personal_info.get('phone'):
        contact_text.append(personal_info['phone'])
    if personal_info.get('location'):
        contact_text.append(personal_info['location'])
    if personal_info.get('linkedin'):
        contact_text.append(personal_info['linkedin'])
    
    contact_para.add_run(' | '.join(contact_text))
    contact_para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Professional Summary
    if summary:
        doc.add_heading('Professional Summary', 2).runs[0].font.color.rgb = RGBColor(0, 51, 102)
        doc.add_paragraph(summary)
        doc.add_paragraph()
    
    # Skills
    if skills:
        doc.add_heading('Skills', 2).runs[0].font.color.rgb = RGBColor(0, 51, 102)
        skills_para = doc.add_paragraph()
        skills_para.add_run(', '.join(skills))
        doc.add_paragraph()
    
    # Work Experience
    if experience:
        doc.add_heading('Work Experience', 2).runs[0].font.color.rgb = RGBColor(0, 51, 102)
        for exp in experience:
            # Job title and company
            job_para = doc.add_paragraph()
            job_run = job_para.add_run(f"{exp.get('title', 'Position')} at {exp.get('company', 'Company')}")
            job_run.bold = True
            
            # Duration
            if exp.get('duration'):
                duration_para = doc.add_paragraph(exp['duration'])
                duration_para.runs[0].italic = True
            
            # Responsibilities
            if exp.get('responsibilities'):
                for resp in exp['responsibilities']:
                    doc.add_paragraph(resp, style='List Bullet')
            
            doc.add_paragraph()
    
    # Education
    if education:
        doc.add_heading('Education', 2).runs[0].font.color.rgb = RGBColor(0, 51, 102)
        for edu in education:
            edu_para = doc.add_paragraph()
            edu_text = f"{edu.get('degree', 'Degree')} in {edu.get('field', 'Field')}"
            if edu.get('institution'):
                edu_text += f" - {edu['institution']}"
            if edu.get('year'):
                edu_text += f" ({edu['year']})"
            edu_para.add_run(edu_text).bold = True
            
            if edu.get('gpa'):
                doc.add_paragraph(f"GPA: {edu['gpa']}")
            doc.add_paragraph()
    
    # Certifications
    if certifications:
        doc.add_heading('Certifications', 2).runs[0].font.color.rgb = RGBColor(0, 51, 102)
        for cert in certifications:
            doc.add_paragraph(cert, style='List Bullet')
    
    # Save to BytesIO
    file_buffer = BytesIO()
    doc.save(file_buffer)
    file_buffer.seek(0)
    
    return file_buffer

def generate_pdf_resume(personal_info, summary, skills, experience, education, certifications):
    """Generate resume in PDF format"""
    from io import BytesIO
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    # Container for elements
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#003366'),
        spaceAfter=12,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#003366'),
        spaceAfter=6,
        borderWidth=1,
        borderColor=colors.HexColor('#003366'),
        borderPadding=3
    )
    
    # Name
    elements.append(Paragraph(personal_info.get('name', 'Your Name'), title_style))
    
    # Contact
    contact_text = []
    if personal_info.get('email'):
        contact_text.append(personal_info['email'])
    if personal_info.get('phone'):
        contact_text.append(personal_info['phone'])
    if personal_info.get('location'):
        contact_text.append(personal_info['location'])
    
    contact_para = Paragraph(' | '.join(contact_text), styles['Normal'])
    contact_para.alignment = 1
    elements.append(contact_para)
    elements.append(Spacer(1, 0.2*inch))
    
    # Summary
    if summary:
        elements.append(Paragraph('Professional Summary', heading_style))
        elements.append(Paragraph(summary, styles['Normal']))
        elements.append(Spacer(1, 0.15*inch))
    
    # Skills
    if skills:
        elements.append(Paragraph('Skills', heading_style))
        elements.append(Paragraph(', '.join(skills), styles['Normal']))
        elements.append(Spacer(1, 0.15*inch))
    
    # Experience
    if experience:
        elements.append(Paragraph('Work Experience', heading_style))
        for exp in experience:
            job_text = f"<b>{exp.get('title', 'Position')} at {exp.get('company', 'Company')}</b>"
            elements.append(Paragraph(job_text, styles['Normal']))
            
            if exp.get('duration'):
                elements.append(Paragraph(f"<i>{exp['duration']}</i>", styles['Normal']))
            
            if exp.get('responsibilities'):
                for resp in exp['responsibilities']:
                    elements.append(Paragraph(f"• {resp}", styles['Normal']))
            
            elements.append(Spacer(1, 0.1*inch))
    
    # Education
    if education:
        elements.append(Paragraph('Education', heading_style))
        for edu in education:
            edu_text = f"<b>{edu.get('degree', 'Degree')} in {edu.get('field', 'Field')}</b>"
            if edu.get('institution'):
                edu_text += f" - {edu['institution']}"
            if edu.get('year'):
                edu_text += f" ({edu['year']})"
            elements.append(Paragraph(edu_text, styles['Normal']))
            elements.append(Spacer(1, 0.1*inch))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    return buffer


if __name__ == '__main__':
    print("\n" + "="*70)
    print("RESUME ANALYZER AND JOB PREDICTOR - ML POWERED")
    print("="*70)
    print("\n[OK] Using Trained ML Model: Gradient Boosting (99.66% Accuracy)")
    print("[OK] Server starting on http://127.0.0.1:5000")
    print("[OK] Job Matching: 21,865 opportunities available")
    print("\n" + "="*70)
    
    app.run(host='0.0.0.0', port=5000, debug=False)
